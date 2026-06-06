"""简历导出服务：生成 PDF / Word / TXT / Markdown 格式的精美排版简历。"""

import re
import io
import os
import logging
from typing import Any

logger = logging.getLogger(__name__)

_FONT_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "fonts", "NotoSansSC-Regular.ttf")

_CJK_FONTS = [
    "/System/Library/Fonts/PingFang.ttc",
    "/www/server/panel/BTPanel/static/font/AlibabaPuHuiTi-2-75-SemiBold-new.ttf",
    "/www/server/panel/BTPanel/static/font/AlibabaPuHuiTi-2-105-Heavy-new.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
    "/www/server/panel/class/fonts/NotoSansSC-Regular.ttf",
    "C:/Windows/Fonts/msyh.ttc",
]

SECTION_KEYWORDS = [
    "求职意向", "个人简介", "自我评价", "个人总结",
    "教育背景", "教育经历", "学历",
    "工作经历", "工作经验",
    "项目经历", "项目经验",
    "专业技能", "技能", "技术栈",
    "证书", "语言能力", "语言",
    "summary", "education", "experience", "skills", "projects", "certifications",
]

CONTACT_PATTERNS = [
    re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+"),
    re.compile(r"1[3-9]\d{9}"),
]


def _parse_sections(text: str) -> dict[str, Any]:
    result: dict[str, Any] = {"name": "", "contact": "", "sections": []}
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return result

    result["name"] = lines[0]

    contact_parts: list[str] = []
    body_start = 1

    for i, line in enumerate(lines[1:], 1):
        is_contact = (
            any(p.search(line) for p in CONTACT_PATTERNS)
            or bool(re.search(r"[|｜·•/]", line))
            or (len(line) < 60 and any(kw in line for kw in ["电话", "邮箱", "手机", "地址", "年龄", "性别", "生日"]))
        )
        is_section_title = any(kw in line for kw in SECTION_KEYWORDS) and len(line) < 20

        if is_section_title:
            body_start = i
            break
        elif is_contact:
            contact_parts.append(line)
            body_start = i + 1
        else:
            body_start = i
            break

    result["contact"] = " | ".join(contact_parts) if contact_parts else ""

    current_section: dict[str, Any] = {"title": "基本信息", "items": []}
    for line in lines[body_start:]:
        is_section_title = any(kw in line for kw in SECTION_KEYWORDS) and len(line) < 20
        if is_section_title:
            if current_section["items"]:
                result["sections"].append(current_section)
            current_section = {"title": line, "items": []}
            continue
        current_section["items"].append(line)
    if current_section["items"]:
        result["sections"].append(current_section)

    return result


def _build_txt(parsed: dict) -> bytes:
    parts: list[str] = []
    parts.append(parsed["name"])
    if parsed["contact"]:
        parts.append(parsed["contact"])
    parts.append("=" * 60)
    for sec in parsed["sections"]:
        parts.append(f"\n【{sec['title']}】")
        parts.append("-" * 40)
        for item in sec["items"]:
            parts.append(f"  ● {item}")
        parts.append("")
    return "\n".join(parts).encode("utf-8")


def _build_md(parsed: dict) -> bytes:
    parts: list[str] = []
    parts.append(f"# {parsed['name']}")
    if parsed["contact"]:
        parts.append(f"\n*{parsed['contact']}*\n")
    parts.append("---\n")
    for sec in parsed["sections"]:
        parts.append(f"## {sec['title']}\n")
        for item in sec["items"]:
            parts.append(f"- {item}")
        parts.append("")
    return "\n".join(parts).encode("utf-8")


def _get_cjk_font_path() -> str:
    for path in _CJK_FONTS:
        if path.endswith('.ttc'):
            continue
        if os.path.exists(path):
            return path
    if os.path.exists(_FONT_PATH):
        return _FONT_PATH
    return ""


def _wrap_text(pdf, text: str, max_width: float) -> list[str]:
    """将文本按指定宽度自动换行，返回行列表。"""
    lines: list[str] = []
    current = ""
    for char in text:
        test = current + char
        if pdf.get_string_width(test) > max_width:
            lines.append(current)
            current = char
        else:
            current = test
    if current:
        lines.append(current)
    return lines if lines else [text]


def _build_pdf(parsed: dict) -> bytes:
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=18)

    font_path = _get_cjk_font_path()
    if font_path:
        pdf.add_font("CJK", "", font_path, uni=True)
        font_name = "CJK"
    else:
        logger.warning("No Chinese font found, falling back to Helvetica")
        font_name = "Helvetica"

    pdf.add_page()

    LEFT = 22
    RIGHT = 188
    PAGE_W = RIGHT - LEFT

    # ── Header bar ──
    pdf.set_fill_color(44, 62, 80)
    pdf.rect(0, 0, 210, 38, "F")

    pdf.set_y(8)
    pdf.set_font(font_name, "", 24)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 10, parsed["name"], ln=1, align="C")

    if parsed["contact"]:
        pdf.set_font(font_name, "", 8.5)
        pdf.set_text_color(200, 210, 220)
        pdf.cell(0, 5, parsed["contact"], ln=1, align="C")

    pdf.set_y(42)
    pdf.set_text_color(30, 30, 30)

    for sec in parsed["sections"]:
        y_before = pdf.get_y()
        if y_before > 255:
            pdf.add_page()
            y_before = pdf.get_y()

        # Section title with accent bar
        pdf.set_fill_color(235, 240, 245)
        pdf.rect(LEFT - 2, pdf.get_y(), 4, 8, "F")
        pdf.set_y(y_before)

        pdf.set_font(font_name, "", 13)
        pdf.set_text_color(44, 62, 80)
        pdf.set_x(LEFT + 6)
        pdf.cell(PAGE_W - 6, 8, sec["title"], ln=1)
        pdf.ln(1)

        # Section divider line
        pdf.set_draw_color(200, 210, 220)
        pdf.set_line_width(0.3)
        pdf.line(LEFT, pdf.get_y(), RIGHT, pdf.get_y())
        pdf.ln(3)

        pdf.set_font(font_name, "", 9.5)
        pdf.set_text_color(55, 55, 55)
        for item in sec["items"]:
            y = pdf.get_y()
            if y > 272:
                pdf.add_page()
            pdf.set_x(LEFT + 2)
            pdf.set_text_color(44, 62, 80)
            pdf.cell(4, 5, ">")
            pdf.set_x(LEFT + 8)
            pdf.set_text_color(55, 55, 55)
            wrapped = _wrap_text(pdf, item, PAGE_W - 12)
            for wi, wline in enumerate(wrapped):
                if wi == 0:
                    pdf.cell(PAGE_W - 12, 5, wline, ln=1)
                else:
                    pdf.set_x(LEFT + 8)
                    pdf.cell(PAGE_W - 12, 5, wline, ln=1)
            pdf.ln(1.5)

        pdf.ln(3)

    return pdf.output(dest="S").encode("latin-1")


def _build_docx(parsed: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Header table with colored background ──
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
    cell._tc.get_or_add_tcPr().append(shading)

    name_para = cell.paragraphs[0]
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(12)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(parsed["name"])
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(255, 255, 255)

    if parsed["contact"]:
        contact_para = cell.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(10)
        contact_run = contact_para.add_run(parsed["contact"])
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(200, 210, 220)

    for row in table.rows:
        for c in row.cells:
            c.width = Inches(6.6)
    doc.add_paragraph()

    for sec in parsed["sections"]:
        heading = doc.add_heading(sec["title"], level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
        for run in heading.runs:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(44, 62, 80)

        pPr = heading._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="2" w:color="8899AA"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        for item in sec["items"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(item)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(60, 60, 60)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_resume(polished_text: str, fmt: str) -> tuple[bytes, str]:
    parsed = _parse_sections(polished_text)

    fmt = fmt.lower()
    if fmt == "pdf":
        content = _build_pdf(parsed)
        return content, "application/pdf"
    elif fmt in ("docx", "word"):
        content = _build_docx(parsed)
        return content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "txt":
        content = _build_txt(parsed)
        return content, "text/plain; charset=utf-8"
    elif fmt == "md":
        content = _build_md(parsed)
        return content, "text/markdown; charset=utf-8"
    else:
        raise ValueError(f"不支持的导出格式: {fmt}")
